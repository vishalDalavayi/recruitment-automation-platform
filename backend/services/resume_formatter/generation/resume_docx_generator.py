import re
from dataclasses import dataclass
from pathlib import Path

from services.resume_formatter.config.settings import Settings
from services.resume_formatter.schemas.resume_schema import Academic, AcademicProject, Experience, ResumeContent


FONT_NAME = "Bookman Old Style"
FONT_SIZE_PT = 10
DEFAULT_BULLET_INDENT_INCHES = 0.5
DOCX_MEDIA_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@dataclass
class GeneratedResumeArtifact:
    file_name: str
    file_path: str


class ResumeDocxGenerator:
    def __init__(self, settings: Settings) -> None:
        self.output_dir = settings.generated_resume_dir.expanduser().resolve()

    def generate(
        self,
        *,
        record_id: str,
        source_file_name: str,
        resume_content: ResumeContent,
    ) -> GeneratedResumeArtifact:
        try:
            from docx import Document
        except ImportError as exc:  # pragma: no cover - depends on runtime environment
            raise RuntimeError("python-docx is required to generate DOCX resumes.") from exc

        self.output_dir.mkdir(parents=True, exist_ok=True)
        file_name = self._build_file_name(
            record_id=record_id,
            source_file_name=source_file_name,
            resume_content=resume_content,
        )
        file_path = self.output_dir / file_name

        document = Document()
        self._configure_document(document)
        self._add_header(document, resume_content)
        self._add_summary(document, resume_content)
        self._add_academics(document, resume_content)
        self._add_technical_skills(document, resume_content)
        self._add_professional_experience(document, resume_content)
        self._add_academic_projects(document, resume_content)
        self._add_certifications(document, resume_content)
        document.save(file_path)

        return GeneratedResumeArtifact(
            file_name=file_name,
            file_path=str(file_path),
        )

    def _configure_document(self, document) -> None:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        from docx.shared import Inches, Pt

        section = document.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        for style_name in ("Normal", "List Paragraph", "No Spacing", "List Bullet"):
            if style_name not in document.styles:
                continue
            style = document.styles[style_name]
            style.font.name = FONT_NAME
            style.font.size = Pt(FONT_SIZE_PT)
            self._set_east_asia_font(style.font, FONT_NAME)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            if style_name == "Normal":
                style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    def _add_header(self, document, resume_content: ResumeContent) -> None:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        if resume_content.Name:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.add_run(resume_content.Name.strip().upper())
            self._format_run(run, bold=True)

        if resume_content.Phone or resume_content.Email:
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

            if resume_content.Phone:
                self._format_run(paragraph.add_run(resume_content.Phone.strip()))

            if resume_content.Phone and resume_content.Email:
                self._format_run(paragraph.add_run(" | "))

            if resume_content.Email:
                self._add_hyperlink(
                    paragraph,
                    text=resume_content.Email.strip(),
                    url=f"mailto:{resume_content.Email.strip()}",
                )

        document.add_paragraph()

    def _add_summary(self, document, resume_content: ResumeContent) -> None:
        summary_paragraphs = self._split_summary_text(resume_content.Summary)
        if not summary_paragraphs:
            return

        self._add_section_heading(document, "SUMMARY")
        for summary_line in summary_paragraphs:
            self._add_bullet_paragraph(document, summary_line, justify=True)

        self._add_blank_paragraph(document)

    def _add_academics(self, document, resume_content: ResumeContent) -> None:
        if not resume_content.Academics:
            return

        self._add_section_heading(document, "EDUCATION")
        for academic in resume_content.Academics:
            self._add_body_paragraph(document, self._format_academic_line(academic), justify=False)

    def _add_technical_skills(self, document, resume_content: ResumeContent) -> None:
        if not resume_content.Technical_Skills:
            return

        document.add_paragraph()
        self._add_section_heading(document, "TECHNICAL SKILLS:")

        for category, skills in resume_content.Technical_Skills.items():
            if not skills:
                continue

            paragraph = self._add_bullet_paragraph(document, "", justify=False)
            self._format_run(paragraph.add_run(category), bold=True)
            self._format_run(paragraph.add_run(f": {', '.join(skill for skill in skills if skill)}"))

    def _add_professional_experience(self, document, resume_content: ResumeContent) -> None:
        experiences = [experience for experience in resume_content.Professional_Experience if self._experience_has_content(experience)]
        if not experiences:
            return

        document.add_paragraph()
        self._add_section_heading(document, "PROFESSIONAL EXPERIENCE")
        document.add_paragraph()

        for index, experience in enumerate(experiences):
            if index > 0:
                document.add_paragraph()

            company_line = self._join_nonempty([experience.Company, experience.location], " — ")
            if company_line:
                paragraph = document.add_paragraph()
                self._set_paragraph_alignment(paragraph, justify=False)
                self._format_run(paragraph.add_run(company_line), bold=True)

            title_dates_line = self._join_nonempty([experience.title, experience.dates_of_employment], " | ")
            if title_dates_line:
                paragraph = document.add_paragraph()
                self._set_paragraph_alignment(paragraph, justify=False)
                self._normalize_paragraph_spacing(paragraph)
                self._format_run(paragraph.add_run(title_dates_line))

            if experience.project_description:
                if title_dates_line:
                    self._add_blank_paragraph(document)
                self._add_body_paragraph(document, experience.project_description, justify=True)

            if experience.Responsibilities:
                if experience.project_description:
                    self._add_blank_paragraph(document)
                paragraph = document.add_paragraph()
                self._set_paragraph_alignment(paragraph, justify=False)
                self._normalize_paragraph_spacing(paragraph)
                self._format_run(paragraph.add_run("Responsibilities:"), bold=True)

                for responsibility in experience.Responsibilities:
                    self._add_bullet_paragraph(document, responsibility, justify=True)

            if experience.Environment:
                document.add_paragraph()
                paragraph = document.add_paragraph()
                self._set_paragraph_alignment(paragraph, justify=False)
                self._normalize_paragraph_spacing(paragraph)
                self._format_run(paragraph.add_run("Environment:"), bold=True)
                self._format_run(paragraph.add_run(f" {', '.join(item for item in experience.Environment if item)}"))

    def _add_academic_projects(self, document, resume_content: ResumeContent) -> None:
        projects = [project for project in resume_content.Academic_projects if self._project_has_content(project)]
        if not projects:
            return

        document.add_paragraph()
        self._add_section_heading(document, "ACADEMIC PROJECTS")
        document.add_paragraph()

        for index, project in enumerate(projects):
            if index > 0:
                document.add_paragraph()

            if project.project_name:
                paragraph = document.add_paragraph()
                self._set_paragraph_alignment(paragraph, justify=False)
                self._normalize_paragraph_spacing(paragraph)
                self._format_run(paragraph.add_run(project.project_name), bold=True)

            if project.your_title:
                self._add_body_paragraph(document, project.your_title, justify=False)

            for responsibility in project.project_responsibilities:
                self._add_body_paragraph(document, responsibility, justify=True)

            if project.Technologies_used:
                document.add_paragraph()
                paragraph = document.add_paragraph()
                self._set_paragraph_alignment(paragraph, justify=False)
                self._normalize_paragraph_spacing(paragraph)
                self._format_run(paragraph.add_run("Technologies:"), bold=True)
                self._format_run(paragraph.add_run(f" {', '.join(item for item in project.Technologies_used if item)}"))

    def _add_certifications(self, document, resume_content: ResumeContent) -> None:
        certifications = [certification.strip() for certification in resume_content.certifications if certification and certification.strip()]
        if not certifications:
            return

        document.add_paragraph()
        self._add_section_heading(document, "CERTIFICATIONS")

        for certification in certifications:
            self._add_bullet_paragraph(document, certification, justify=False)

    def _add_section_heading(self, document, heading: str):
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        self._normalize_paragraph_spacing(paragraph)
        run = paragraph.add_run(heading)
        self._format_run(run, bold=True)
        return paragraph

    def _add_body_paragraph(self, document, text: str, *, justify: bool, style_name: str | None = None):
        paragraph = document.add_paragraph(style=style_name)
        self._set_paragraph_alignment(paragraph, justify=justify)
        self._normalize_paragraph_spacing(paragraph)
        self._format_run(paragraph.add_run(text.strip()))
        return paragraph

    def _add_bullet_paragraph(self, document, text: str, *, justify: bool):
        style_name = "List Bullet" if "List Bullet" in document.styles else None
        paragraph = document.add_paragraph(style=style_name)
        self._set_paragraph_alignment(paragraph, justify=justify)
        self._normalize_paragraph_spacing(paragraph)
        paragraph.paragraph_format.left_indent = self._inches(DEFAULT_BULLET_INDENT_INCHES)
        if text:
            self._format_run(paragraph.add_run(text.strip()))
        return paragraph

    def _add_blank_paragraph(self, document):
        paragraph = document.add_paragraph()
        self._normalize_paragraph_spacing(paragraph)
        return paragraph

    def _set_paragraph_alignment(self, paragraph, *, justify: bool) -> None:
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY if justify else WD_PARAGRAPH_ALIGNMENT.LEFT

    def _format_run(self, run, *, bold: bool = False) -> None:
        run.font.name = FONT_NAME
        run.font.size = self._pt(FONT_SIZE_PT)
        run.bold = bold
        self._set_east_asia_font(run.font, FONT_NAME)

    def _normalize_paragraph_spacing(self, paragraph) -> None:
        paragraph.paragraph_format.space_before = self._pt(0)
        paragraph.paragraph_format.space_after = self._pt(0)

    def _add_hyperlink(self, paragraph, *, text: str, url: str) -> None:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        part = paragraph.part
        relationship_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), relationship_id)

        run = OxmlElement("w:r")
        run_properties = OxmlElement("w:rPr")

        r_fonts = OxmlElement("w:rFonts")
        r_fonts.set(qn("w:ascii"), FONT_NAME)
        r_fonts.set(qn("w:hAnsi"), FONT_NAME)
        r_fonts.set(qn("w:eastAsia"), FONT_NAME)
        run_properties.append(r_fonts)

        bold = OxmlElement("w:b")
        bold.set(qn("w:val"), "1")
        run_properties.append(bold)

        color = OxmlElement("w:color")
        color.set(qn("w:val"), "1155CC")
        run_properties.append(color)

        underline = OxmlElement("w:u")
        underline.set(qn("w:val"), "single")
        run_properties.append(underline)

        size = OxmlElement("w:sz")
        size.set(qn("w:val"), str(FONT_SIZE_PT * 2))
        run_properties.append(size)

        size_cs = OxmlElement("w:szCs")
        size_cs.set(qn("w:val"), str(FONT_SIZE_PT * 2))
        run_properties.append(size_cs)

        run.append(run_properties)
        text_element = OxmlElement("w:t")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)
        paragraph._p.append(hyperlink)

    def _format_academic_line(self, academic: Academic) -> str:
        degree_major = self._join_nonempty([academic.Degree, academic.Major], " – ")
        if degree_major and academic.University:
            return f"{degree_major} from {academic.University}"
        return degree_major or academic.University

    def _build_file_name(self, *, record_id: str, source_file_name: str, resume_content: ResumeContent) -> str:
        base_name = resume_content.Name.strip() or Path(source_file_name).stem or "generated_resume"
        safe_name = re.sub(r"[^A-Za-z0-9]+", "_", base_name).strip("_").lower() or "generated_resume"
        return f"{safe_name}_{record_id[:8]}.docx"

    def _split_summary_text(self, summary: str) -> list[str]:
        cleaned = summary.strip()
        if not cleaned:
            return []
        if "\n" in cleaned:
            return [line.strip() for line in cleaned.splitlines() if line.strip()]

        parts = re.split(r"(?<=[.!?])\s+(?=[A-Z])", cleaned)
        return [part.strip() for part in parts if part.strip()]

    def _join_nonempty(self, values: list[str], separator: str) -> str:
        return separator.join(value.strip() for value in values if value and value.strip())

    def _experience_has_content(self, experience: Experience) -> bool:
        return any(
            [
                experience.Company,
                experience.location,
                experience.title,
                experience.dates_of_employment,
                experience.project_description,
                experience.Responsibilities,
                experience.Environment,
            ]
        )

    def _project_has_content(self, project: AcademicProject) -> bool:
        return any(
            [
                project.project_name,
                project.your_title,
                project.project_responsibilities,
                project.Technologies_used,
            ]
        )

    def _set_east_asia_font(self, font, font_name: str) -> None:
        from docx.oxml.ns import qn

        font.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _pt(self, value: float):
        from docx.shared import Pt

        return Pt(value)

    def _inches(self, value: float):
        from docx.shared import Inches

        return Inches(value)
