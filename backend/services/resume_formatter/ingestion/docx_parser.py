from io import BytesIO


def parse_docx_bytes(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - depends on runtime environment
        raise RuntimeError("python-docx is required to parse DOCX resumes.") from exc

    try:
        document = Document(BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError("Failed to read the uploaded DOCX file.") from exc

    parts: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text and paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()
